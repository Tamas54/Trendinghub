"""
TrendMaster - Flask Application
Trending topics collector and Facebook post generator
"""
from flask import Flask, render_template, jsonify, request
from flask_cors import CORS
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.interval import IntervalTrigger
from apscheduler.triggers.cron import CronTrigger
from datetime import datetime
import os
import atexit
import logging
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Import local modules
from database import db
from collector import TrendCollector
from news_collector import NewsCollector
from generator import PostGenerator
from google_ai import GoogleAIGenerator
from publisher import SocialPublisher
from super_trends import detector
from media_spoofer import MediaSpoofer
from facebook_poster import publish_to_facebook_sync

# === ÚJ IMPORTS - SaaS rendszer ===
from agent_api import agent_api
from seo_api import seo_api

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'trendmaster-secret-key-2025')

# CORS engedélyezése (Chrome extension támogatás)
CORS(app, origins=['chrome-extension://*', 'http://localhost:*', 'https://*.railway.app'])

# === ÚJ BLUEPRINTS - SaaS rendszer ===
app.register_blueprint(agent_api)  # /api/agent/* végpontok
app.register_blueprint(seo_api)    # /api/optimize-seo, /api/generate-hashtags, stb.

# Initialize components
trend_collector = TrendCollector()
news_collector = NewsCollector()
post_generator = PostGenerator()  # OpenAI
google_ai_generator = GoogleAIGenerator()  # Google AI
social_publisher = SocialPublisher()
media_spoofer = MediaSpoofer()

# AI Provider selection
AI_PROVIDER = os.getenv('AI_PROVIDER', 'openai').lower()

# Configure APScheduler logging
logging.basicConfig()
logging.getLogger('apscheduler').setLevel(logging.INFO)

# Scheduler for automatic trend collection
scheduler = BackgroundScheduler({
    'apscheduler.executors.default': {
        'class': 'apscheduler.executors.pool:ThreadPoolExecutor',
        'max_workers': '1'
    },
    'apscheduler.job_defaults.coalesce': 'true',
    'apscheduler.job_defaults.max_instances': '1'
})


def collect_trends_job():
    """
    Background job to collect trends
    Runs every 12 hours
    """
    print(f"\n{'='*60}")
    print(f"⏰ SCHEDULED TREND COLLECTION")
    print(f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"{'='*60}")

    try:
        # Collect from all sources
        all_trends = trend_collector.get_flat_trends_list()

        # Collect news articles
        news_articles = news_collector.collect_all_news(max_per_source=5)

        # Separate Google News trending from regular news
        google_news_hu = []
        google_news_gb = []
        google_news_us = []
        regular_news = []

        for article in news_articles:
            # Save to database
            db.save_news_article(article)

            # Categorize by source
            if article['source'] == 'Google News HU':
                google_news_hu.append({
                    'source': 'google_hu',
                    'topic': article['title'],
                    'rank': len(google_news_hu) + 1,
                    'relevance_score': 10 - len(google_news_hu),
                    'metadata': f"Category: {article['category']} | Link: {article['link']}"
                })
            elif article['source'] == 'Google News GB':
                google_news_gb.append({
                    'source': 'google_gb',
                    'topic': article['title'],
                    'rank': len(google_news_gb) + 1,
                    'relevance_score': 10 - len(google_news_gb),
                    'metadata': f"Category: {article['category']} | Link: {article['link']}"
                })
            elif article['source'] == 'Google News US':
                google_news_us.append({
                    'source': 'google_us',
                    'topic': article['title'],
                    'rank': len(google_news_us) + 1,
                    'relevance_score': 10 - len(google_news_us),
                    'metadata': f"Category: {article['category']} | Link: {article['link']}"
                })
            else:
                regular_news.append(article)

        # Add Google News trends to main trends
        all_trends.extend(google_news_hu[:10])
        all_trends.extend(google_news_gb[:10])
        all_trends.extend(google_news_us[:10])

        # Extract regular news trends
        news_trends = news_collector.get_trending_topics_from_news(regular_news)
        all_trends.extend(news_trends)

        # Save all trends to database
        saved = db.save_trends(all_trends)

        print(f"✅ Scheduled collection complete: {saved} new trends saved")

    except Exception as e:
        print(f"❌ Scheduled collection error: {e}")


def publish_scheduled_posts_job():
    """
    Background job to publish scheduled posts
    Runs every minute to check for posts that need to be published
    """
    try:
        # Get pending posts that should be published now
        pending_posts = db.get_pending_scheduled_posts()

        if not pending_posts:
            return

        print(f"\n{'='*60}")
        print(f"📅 CHECKING SCHEDULED POSTS")
        print(f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"Found {len(pending_posts)} post(s) to publish")
        print(f"{'='*60}")

        for post in pending_posts:
            post_id = post['id']
            post_content = post['post_content']
            image_path = post.get('image_path')
            video_path = post.get('video_path')
            platform = post.get('platform', 'facebook')

            print(f"\n📤 Publishing scheduled post #{post_id}...")
            print(f"   Platform: {platform}")
            print(f"   Content: {post_content[:50]}...")

            try:
                # Publish based on platform
                if platform == 'facebook':
                    # Use Firefox session-based Facebook poster (no OAuth needed!)
                    result = publish_to_facebook_sync(post_content, image_path)
                else:
                    error_msg = f'Unsupported platform: {platform}'
                    print(f"   ❌ {error_msg}")
                    db.update_scheduled_post_status(
                        post_id=post_id,
                        status='failed',
                        error_message=error_msg
                    )
                    continue

                # Update status based on result
                if result.get('success'):
                    print(f"   ✅ Published successfully!")
                    print(f"   📸 Screenshot: {result.get('screenshot', 'N/A')}")
                    db.update_scheduled_post_status(
                        post_id=post_id,
                        status='published',
                        published_at=datetime.now().isoformat()
                    )
                else:
                    error_msg = result.get('message', 'Unknown error')
                    print(f"   ❌ Publishing failed: {error_msg}")
                    db.update_scheduled_post_status(
                        post_id=post_id,
                        status='failed',
                        error_message=error_msg
                    )

            except Exception as e:
                error_msg = str(e)
                print(f"   ❌ Exception during publishing: {error_msg}")
                db.update_scheduled_post_status(
                    post_id=post_id,
                    status='failed',
                    error_message=error_msg
                )

        print(f"{'='*60}\n")

    except Exception as e:
        print(f"❌ Scheduled posts publishing error: {e}")


# Schedule trend collection every 12 hours using CronTrigger (more reliable)
# Runs at 00:00 and 12:00 every day
scheduler.add_job(
    func=collect_trends_job,
    trigger=CronTrigger(hour='0,12', minute=0),
    id='collect_trends_cron',
    name='Collect trends at midnight and noon',
    replace_existing=True
)

# ALSO add IntervalTrigger as backup (runs every 12 hours from start)
scheduler.add_job(
    func=collect_trends_job,
    trigger=IntervalTrigger(hours=12),
    id='collect_trends_interval',
    name='Collect trends every 12 hours (interval)',
    replace_existing=True
)

# Also run on startup
scheduler.add_job(
    func=collect_trends_job,
    trigger='date',
    run_date=datetime.now(),
    id='initial_collection',
    name='Initial trend collection'
)

# Schedule post publishing job - runs every minute to check for scheduled posts
scheduler.add_job(
    func=publish_scheduled_posts_job,
    trigger=IntervalTrigger(minutes=1),
    id='publish_scheduled_posts',
    name='Publish scheduled posts every minute',
    replace_existing=True
)

# Start scheduler
scheduler.start()

# Print scheduled jobs
print("\n" + "="*60)
print("✅ APScheduler started")
print("="*60)
print("Scheduled jobs:")
for job in scheduler.get_jobs():
    print(f"  • {job.name}")
    print(f"    Next run: {job.next_run_time}")
print("="*60 + "\n")

# Shutdown scheduler on exit
atexit.register(lambda: scheduler.shutdown())


# ============================================================================
# ROUTES
# ============================================================================

@app.route('/')
def index():
    """Landing page"""
    return render_template('landing.html')


@app.route('/dashboard')
def dashboard():
    """Dashboard page with trends"""
    stats = db.get_stats()
    return render_template('dashboard.html', stats=stats)


@app.route('/editor')
def editor():
    """Editor and publishing interface"""
    return render_template('editor.html')


@app.route('/test_news_button')
def test_news_button():
    """Test page for news button onclick"""
    return render_template('test_news_button.html')


@app.route('/health')
def health():
    """Health check endpoint for Railway"""
    stats = db.get_stats()
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'database': stats
    })


@app.route('/api/trends')
def get_trends():
    """
    Get latest trends grouped by source
    Returns top 10 trends for each of 6 sources
    """
    sources = [
        'google_hu', 'google_gb', 'google_us',
        'youtube_hu', 'youtube_gb', 'youtube_us'
    ]

    trends_by_source = {}

    for source in sources:
        trends = db.get_latest_trends(source=source, limit=10)
        trends_by_source[source] = trends

    # Also get news trends
    news_articles = db.get_latest_news(limit=10)
    trends_by_source['news'] = news_articles

    stats = db.get_stats()

    return jsonify({
        'trends': trends_by_source,
        'stats': stats,
        'timestamp': datetime.now().isoformat()
    })


@app.route('/api/super-trends')
def get_super_trends():
    """
    Get super trends - topics that appear across multiple sources
    """
    sources = [
        'google_hu', 'google_gb', 'google_us',
        'youtube_hu', 'youtube_gb', 'youtube_us'
    ]

    trends_by_source = {}

    # Get trends from each source
    for source in sources:
        trends = db.get_latest_trends(source=source, limit=15)
        trends_by_source[source] = trends

    # ALSO get news articles (this was missing!)
    news_articles = db.get_latest_news(limit=30)
    if news_articles:
        # Group news by source
        news_by_source = {}
        for article in news_articles:
            source_name = article.get('source', 'news').lower().replace(' ', '_')
            if source_name not in news_by_source:
                news_by_source[source_name] = []
            news_by_source[source_name].append(article)

        # Add news to trends_by_source
        trends_by_source.update(news_by_source)

    # Detect super trends (topics appearing in 3+ sources)
    super_trends = detector.detect_super_trends(
        trends_by_source,
        min_sources=3,
        similarity_threshold=0.25  # Lowered threshold to catch more matches
    )

    return jsonify({
        'super_trends': super_trends,
        'count': len(super_trends),
        'timestamp': datetime.now().isoformat()
    })


@app.route('/api/trends/<int:trend_id>')
def get_trend_detail(trend_id):
    """Get single trend with any existing generated posts"""
    trend = db.get_trend_by_id(trend_id)

    if not trend:
        return jsonify({'error': 'Trend not found'}), 404

    # Get existing posts
    posts = db.get_posts_for_trend(trend_id)

    return jsonify({
        'trend': trend,
        'posts': posts
    })


@app.route('/api/recent-posts')
def get_recent_posts():
    """Get recently generated posts with their trend info"""
    conn = db.get_connection()
    cursor = conn.cursor()

    # Get last 20 generated posts with trend info
    cursor.execute('''
        SELECT
            p.id, p.trend_id, p.post_text, p.char_count, p.generated_at,
            t.topic, t.source, t.metadata
        FROM generated_posts p
        JOIN trends t ON p.trend_id = t.id
        ORDER BY p.generated_at DESC
        LIMIT 20
    ''')

    rows = cursor.fetchall()
    conn.close()

    posts = [dict(row) for row in rows]

    return jsonify({
        'posts': posts,
        'count': len(posts)
    })


@app.route('/api/news')
def get_news():
    """Get recent news articles"""
    conn = db.get_connection()
    cursor = conn.cursor()

    cursor.execute('''
        SELECT id, source, title, description, link, pub_date, category
        FROM news_articles
        ORDER BY pub_date DESC
        LIMIT 50
    ''')

    rows = cursor.fetchall()
    conn.close()

    news = [dict(row) for row in rows]

    return jsonify({
        'news': news,
        'count': len(news)
    })


@app.route('/api/generate', methods=['POST'])
def generate_posts():
    """
    Generate Facebook posts for a trend or news
    POST body: { "trend_id": 123 } or { "news_id": "abc123" }
    """
    data = request.get_json()
    trend_id = data.get('trend_id')
    news_id = data.get('news_id')

    if not trend_id and not news_id:
        return jsonify({'error': 'trend_id or news_id is required'}), 400

    # Handle trend
    if trend_id:
        trend = db.get_trend_by_id(trend_id)
        if not trend:
            return jsonify({'error': 'Trend not found'}), 404

        # Check if posts already exist
        existing_posts = db.get_posts_for_trend(trend_id)
        if existing_posts:
            return jsonify({
                'trend_id': trend_id,
                'posts': [p['post_text'] for p in existing_posts],
                'source_url': None,  # Trends don't have source URLs
                'cached': True
            })

        topic = trend.get('topic', '')
        source = trend.get('source', '')
        metadata = trend.get('metadata', '')
        source_url = None  # Trends don't have source URLs

    # Handle news
    else:
        # Get news from database
        conn = db.get_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM news_articles WHERE id = ?', (news_id,))
        news = cursor.fetchone()
        conn.close()

        if not news:
            return jsonify({'error': 'News not found'}), 404

        news = dict(news)

        # Check if posts already exist for this news
        conn = db.get_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM generated_posts WHERE trend_id = ?', (news_id,))
        existing = cursor.fetchall()
        conn.close()

        if existing:
            return jsonify({
                'news_id': news_id,
                'posts': [dict(p)['post_text'] for p in existing],
                'source_url': news.get('link'),  # Article link kommenthez
                'cached': True
            })

        topic = news.get('title', '')
        source = f"news_{news.get('source', 'unknown').lower().replace(' ', '_')}"
        # Pass FULL description (up to 2000 chars) for better AI context
        description = news.get('description', '') or ''
        metadata = f"CIKK RÉSZLETES LEÍRÁSA:\n{description[:2000]}\n\nKategória: {news.get('category', '')}"
        trend_id = news_id  # Use news_id as trend_id for storage
        source_url = news.get('link')  # Article link kommenthez

    # Generate new posts
    print(f"🤖 Generating new posts: {topic[:50]}...")
    print(f"   Using AI Provider: {AI_PROVIDER.upper()}")

    try:
        # Choose generator based on AI_PROVIDER
        if AI_PROVIDER == 'google':
            posts = google_ai_generator.generate_facebook_posts(topic, source, metadata)
        else:
            posts = post_generator.generate_facebook_posts(topic, source, metadata)

        # Save posts to database
        for post_text in posts:
            db.save_generated_post(trend_id, post_text)

        return jsonify({
            'trend_id': trend_id,
            'posts': posts,
            'source_url': source_url,  # Include source URL for comments
            'cached': False
        })

    except Exception as e:
        print(f"❌ Error generating posts: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/refresh', methods=['POST'])
def manual_refresh():
    """
    Manually trigger trend collection
    """
    try:
        collect_trends_job()
        return jsonify({
            'success': True,
            'message': 'Trends collected successfully',
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/stats')
def get_stats():
    """Get database statistics"""
    stats = db.get_stats()
    return jsonify(stats)


@app.route('/api/posts/<int:post_id>', methods=['DELETE'])
def delete_post(post_id):
    """Delete a generated post"""
    try:
        conn = db.get_connection()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM generated_posts WHERE id = ?', (post_id,))
        deleted = cursor.rowcount
        conn.commit()
        conn.close()

        if deleted == 0:
            return jsonify({'error': 'Post not found'}), 404

        return jsonify({
            'success': True,
            'message': f'Post {post_id} deleted'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/search')
def search():
    """
    Search trends and posts
    Query params: q=search_term
    """
    query = request.args.get('q', '').strip()

    if not query or len(query) < 2:
        return jsonify({'error': 'Query must be at least 2 characters'}), 400

    conn = db.get_connection()
    cursor = conn.cursor()

    # Search in trends
    cursor.execute('''
        SELECT * FROM trends
        WHERE topic LIKE ? OR source LIKE ? OR metadata LIKE ?
        ORDER BY fetch_time DESC
        LIMIT 20
    ''', (f'%{query}%', f'%{query}%', f'%{query}%'))
    trends = [dict(row) for row in cursor.fetchall()]

    # Search in generated posts
    cursor.execute('''
        SELECT p.*, t.topic, t.source
        FROM generated_posts p
        LEFT JOIN trends t ON p.trend_id = t.id
        WHERE p.post_text LIKE ?
        ORDER BY p.generated_at DESC
        LIMIT 20
    ''', (f'%{query}%',))
    posts = [dict(row) for row in cursor.fetchall()]

    # Search in news
    cursor.execute('''
        SELECT * FROM news_articles
        WHERE title LIKE ? OR description LIKE ? OR category LIKE ?
        ORDER BY fetch_time DESC
        LIMIT 20
    ''', (f'%{query}%', f'%{query}%', f'%{query}%'))
    news = [dict(row) for row in cursor.fetchall()]

    conn.close()

    return jsonify({
        'query': query,
        'trends': trends,
        'posts': posts,
        'news': news,
        'total': len(trends) + len(posts) + len(news)
    })


@app.route('/api/cleanup', methods=['POST'])
def cleanup_old_data():
    """Clean up old trends (older than 7 days)"""
    data = request.get_json() or {}
    days = data.get('days', 7)

    try:
        deleted = db.cleanup_old_data(days=days)
        return jsonify({
            'success': True,
            'deleted': deleted,
            'message': f'Deleted {deleted} old trends'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/save-connection', methods=['POST'])
def save_connection():
    """
    Save Nango connection ID for social media
    POST body: { "provider": "facebook", "connection_id": "conn_xyz", "profile_name": "optional" }
    """
    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    provider = data.get('provider')
    connection_id = data.get('connection_id')
    profile_name = data.get('profile_name')

    if not provider or not connection_id:
        return jsonify({'error': 'provider and connection_id are required'}), 400

    try:
        success = db.save_connection(provider, connection_id, profile_name)

        if success:
            return jsonify({
                'success': True,
                'message': f'Connection saved for {provider}',
                'provider': provider
            })
        else:
            return jsonify({'error': 'Failed to save connection'}), 500

    except Exception as e:
        print(f"❌ Error saving connection: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-image-prompt', methods=['POST'])
def generate_image_prompt():
    """
    Generate an optimized image prompt from post text using GPT-5
    POST body: { "post_text": "the post content" }
    Returns: { "success": true, "prompt": "optimized image prompt" }
    """
    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    post_text = data.get('post_text')

    if not post_text:
        return jsonify({'error': 'post_text is required'}), 400

    try:
        print(f"📝 Generating image prompt from post text...")
        prompt = post_generator.generate_image_prompt(post_text)

        return jsonify({
            'success': True,
            'prompt': prompt
        })

    except Exception as e:
        print(f"❌ Error generating image prompt: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-image', methods=['POST'])
def generate_image():
    """
    Generate image using Nano Banana (Gemini 3 Pro Image) - default
    POST body: {
        "prompt": "description of image",
        "provider": "google" | "openai" (optional, default: google for Nano Banana)
    }
    """
    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    prompt = data.get('prompt')
    # Default to Google (Nano Banana / Gemini 3 Pro Image)
    provider = data.get('provider', 'google')

    if not prompt:
        return jsonify({'error': 'prompt is required'}), 400

    try:
        print(f"🎨 Generating image with Nano Banana (Gemini 3 Pro Image)")

        # Always use Google/Nano Banana for image generation
        result = google_ai_generator.generate_image(prompt)

        # Check if result is a local file path or URL
        if result and result.startswith('/') and not result.startswith('http'):
            # Local file path - convert to served URL
            image_url = f"/api/serve-image{result}"
        else:
            # Already a URL (fallback case)
            image_url = result

        return jsonify({
            'success': True,
            'image_url': image_url,
            'prompt': prompt,
            'provider': 'google (Nano Banana)'
        })

    except Exception as e:
        print(f"❌ Error generating image: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-video-prompt', methods=['POST'])
def generate_video_prompt():
    """
    Generate a clean video prompt from Facebook post text using Gemini AI
    POST body: { "post_text": "Facebook post content..." }
    """
    try:
        data = request.get_json()

        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        post_text = data.get('post_text', '')

        if not post_text:
            return jsonify({'error': 'post_text is required'}), 400

        print(f"🤖 Generating video prompt from post: {post_text[:50]}...")

        # Use Google AI to generate clean prompt
        video_prompt = google_ai_generator.generate_video_prompt_from_post(post_text)

        return jsonify({
            'success': True,
            'video_prompt': video_prompt
        })

    except Exception as e:
        print(f"❌ Error generating video prompt: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-video', methods=['POST'])
def generate_video():
    """
    Generate video using Veo 2.0 (Google AI only)
    POST body: { "prompt": "description of video", "duration": 5 }
    """
    try:
        print("🔧 DEBUG: generate_video endpoint called")
        data = request.get_json()
        print(f"🔧 DEBUG: data = {data}")

        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        prompt = data.get('prompt')
        duration = data.get('duration', 5)  # Default 5 seconds
        print(f"🔧 DEBUG: prompt = {prompt}, duration = {duration}")

        if not prompt:
            return jsonify({'error': 'prompt is required'}), 400
    except Exception as e:
        print(f"❌ Error in request parsing: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Request parsing failed: {str(e)}'}), 500

    try:
        print(f"🎬 Generating video: {prompt[:50]}...")
        print(f"   Using AI Provider for video: GOOGLE (Veo 3.1)")

        # Always use Google Veo 3.1 for video generation
        # (Sora requires verified organization, so we always use Google for video)
        video_path = google_ai_generator.generate_video(prompt, duration)

        if not video_path:
            return jsonify({'error': 'Video generation failed'}), 500

        # Return the video file path
        # We'll need to serve this file via a separate endpoint
        return jsonify({
            'success': True,
            'video_path': video_path,
            'prompt': prompt,
            'duration': duration
        })

    except Exception as e:
        print(f"❌ Error generating video: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/serve-video/<path:filename>', methods=['GET'])
def serve_video(filename):
    """Serve generated video file"""
    try:
        from flask import send_file
        import os

        print(f"🎥 Serving video: {filename}")

        # Ensure absolute path (add leading / if missing)
        if not filename.startswith('/'):
            filename = '/' + filename

        print(f"🎥 Absolute path: {filename}")

        # Check if file exists
        if os.path.exists(filename):
            print(f"✅ Video file found, serving...")
            return send_file(filename, mimetype='video/mp4')
        else:
            print(f"❌ Video file not found: {filename}")
            return jsonify({'error': 'Video file not found'}), 404
    except Exception as e:
        print(f"❌ Error serving video: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/serve-image/<path:filename>', methods=['GET'])
def serve_image(filename):
    """Serve generated image file (Nano Banana)"""
    try:
        from flask import send_file
        import os

        print(f"🖼️ Serving image: {filename}")

        # Ensure absolute path (add leading / if missing)
        if not filename.startswith('/'):
            filename = '/' + filename

        print(f"🖼️ Absolute path: {filename}")

        # Check if file exists
        if os.path.exists(filename):
            print(f"✅ Image file found, serving...")
            return send_file(filename, mimetype='image/png')
        else:
            print(f"❌ Image file not found: {filename}")
            return jsonify({'error': 'Image file not found'}), 404
    except Exception as e:
        print(f"❌ Error serving image: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/publish', methods=['POST'])
def publish_post():
    """
    Publish post to social media
    POST body: {
        "message": "post content",
        "provider": "facebook",
        "image_url": "optional",
        "source_url": "optional - eredeti cikk link"
    }
    """
    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    message = data.get('message')
    provider = data.get('provider')
    image_url = data.get('image_url')
    source_url = data.get('source_url')  # Eredeti link kommenthez

    print(f"🔍 DEBUG - source_url érték: '{source_url}'")
    print(f"🔍 DEBUG - source_url típus: {type(source_url)}")

    if not message or not provider:
        return jsonify({'error': 'message and provider are required'}), 400

    image_path = None
    try:
        # Ha van image_url, letöltjük
        if image_url:
            import requests
            import tempfile
            print(f"📥 Letöltés: {image_url}")

            response = requests.get(image_url, timeout=10)
            if response.status_code == 200:
                # Temp file létrehozása
                suffix = '.jpg'
                if 'image/png' in response.headers.get('content-type', ''):
                    suffix = '.png'

                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                temp_file.write(response.content)
                temp_file.close()
                image_path = temp_file.name
                print(f"✅ Kép letöltve: {image_path}")
            else:
                print(f"⚠️  Image download failed: {response.status_code}")

        # Publish based on provider
        if provider == 'facebook':
            # Use Firefox session-based Facebook poster (no OAuth needed!)
            result = publish_to_facebook_sync(
                message,
                image_path,
                comment_text=source_url  # Forrás link kommentként
            )
        else:
            return jsonify({'error': f'Unsupported provider: {provider}'}), 400

        # Cleanup temp file
        if image_path and os.path.exists(image_path):
            try:
                os.remove(image_path)
            except:
                pass

        if result.get('success'):
            return jsonify({
                'success': True,
                'message': result.get('message', 'Post published successfully'),
                'screenshot': result.get('screenshot'),
                'provider': provider
            })
        else:
            return jsonify({
                'success': False,
                'error': result.get('message', 'Unknown error')
            }), 500

    except Exception as e:
        print(f"❌ Error publishing post: {e}")
        import traceback
        traceback.print_exc()

        # Cleanup temp file on error
        if image_path and os.path.exists(image_path):
            try:
                os.remove(image_path)
            except:
                pass

        return jsonify({'error': str(e)}), 500


@app.route('/api/download-image', methods=['POST'])
def download_image():
    """
    Download an image from URL (proxy to avoid CORS)
    POST body: { "image_url": "https://..." or "/api/serve-image/..." }
    """
    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    image_url = data.get('image_url')

    if not image_url:
        return jsonify({'error': 'image_url is required'}), 400

    try:
        import tempfile
        import uuid
        from flask import send_file

        print(f"📥 Downloading image from: {image_url[:80]}...")

        # Check if it's a local serve-image URL
        if image_url.startswith('/api/serve-image/'):
            # Extract the file path from the URL
            file_path = image_url.replace('/api/serve-image', '')
            if not file_path.startswith('/'):
                file_path = '/' + file_path

            print(f"📥 Local image path: {file_path}")

            if os.path.exists(file_path):
                return send_file(
                    file_path,
                    mimetype='image/png',
                    as_attachment=True,
                    download_name=f"trendmaster-{datetime.now().strftime('%Y%m%d-%H%M%S')}.png"
                )
            else:
                return jsonify({'error': f'Local file not found: {file_path}'}), 404
        else:
            # External URL - download via requests
            import requests as req

            response = req.get(image_url, timeout=10)
            response.raise_for_status()

            # Save to temp file
            temp_dir = tempfile.gettempdir()
            temp_filename = f"download_{uuid.uuid4()}.png"
            temp_path = os.path.join(temp_dir, temp_filename)

            with open(temp_path, 'wb') as f:
                f.write(response.content)

            print(f"✅ Image downloaded successfully")

            return send_file(
                temp_path,
                mimetype='image/png',
                as_attachment=True,
                download_name=f"trendmaster-{datetime.now().strftime('%Y%m%d-%H%M%S')}.png"
            )

    except Exception as e:
        print(f"❌ Error downloading image: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/spoof-image', methods=['POST'])
def spoof_image():
    """
    Apply EXIF spoofing to an uploaded image
    Accepts multipart/form-data with:
    - image: file upload
    - device: optional device key (default: random)
    """
    if 'image' not in request.files:
        return jsonify({'error': 'No image file provided'}), 400

    file = request.files['image']
    device = request.form.get('device', 'random')

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        # Save uploaded file temporarily
        import tempfile
        import uuid
        from PIL import Image as PILImage

        temp_dir = tempfile.gettempdir()

        # Always save as .png first to let spoofer detect and convert
        temp_filename = f"spoof_{uuid.uuid4()}.png"
        temp_path = os.path.join(temp_dir, temp_filename)

        file.save(temp_path)
        print(f"🔧 Saved uploaded file: {temp_path}")
        print(f"🔧 File size: {os.path.getsize(temp_path)} bytes")

        # Check actual image format
        try:
            with PILImage.open(temp_path) as img:
                print(f"🔧 Image format: {img.format}, mode: {img.mode}, size: {img.size}")
        except Exception as img_err:
            print(f"⚠️ Could not read image info: {img_err}")

        # Apply EXIF spoofing
        print(f"🔧 Spoofing image: {temp_path} with device: {device}")
        success = media_spoofer.spoof_photo(temp_path, device_key=device)
        print(f"🔧 Spoof result: {success}")

        if not success:
            print(f"❌ Spoof failed for: {temp_path}")
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return jsonify({'error': 'Failed to spoof image'}), 500

        # Return the spoofed image
        from flask import send_file
        return send_file(
            temp_path,
            mimetype='image/jpeg',
            as_attachment=True,
            download_name=f"spoofed_{file.filename}"
        )

    except Exception as e:
        print(f"❌ Error spoofing image: {e}")
        import traceback
        traceback.print_exc()
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({'error': str(e)}), 500


@app.route('/api/spoof-video', methods=['POST'])
def spoof_video():
    """
    Apply metadata spoofing to an uploaded video
    Accepts multipart/form-data with:
    - video: file upload
    - device: optional device key (default: random)
    """
    if 'video' not in request.files:
        return jsonify({'error': 'No video file provided'}), 400

    file = request.files['video']
    device = request.form.get('device', 'random')

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        # Save uploaded file temporarily
        import tempfile
        import uuid

        temp_dir = tempfile.gettempdir()
        file_ext = os.path.splitext(file.filename)[1] or '.mp4'
        temp_filename = f"spoof_{uuid.uuid4()}{file_ext}"
        temp_path = os.path.join(temp_dir, temp_filename)

        file.save(temp_path)

        # Apply video metadata spoofing
        success = media_spoofer.spoof_video(temp_path, device_key=device)

        if not success:
            os.remove(temp_path)
            return jsonify({'error': 'Failed to spoof video'}), 500

        # Return the spoofed video
        from flask import send_file
        return send_file(
            temp_path,
            mimetype='video/mp4',
            as_attachment=True,
            download_name=f"spoofed_{file.filename}"
        )

    except Exception as e:
        print(f"❌ Error spoofing video: {e}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({'error': str(e)}), 500


@app.route('/api/connection-status')
def connection_status():
    """
    Check if social media connection exists
    Query params: provider (e.g., 'facebook')
    """
    provider = request.args.get('provider', 'facebook')

    try:
        connection_id = db.get_connection_id(provider)

        if connection_id:
            # Get profile name if available
            conn = db.get_connection()
            cursor = conn.cursor()
            cursor.execute('SELECT profile_name FROM social_connections WHERE provider = ?', (provider,))
            row = cursor.fetchone()
            conn.close()

            return jsonify({
                'connected': True,
                'provider': provider,
                'profile_name': row['profile_name'] if row and row['profile_name'] else None
            })
        else:
            return jsonify({
                'connected': False,
                'provider': provider
            })

    except Exception as e:
        print(f"❌ Error checking connection status: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/disconnect', methods=['POST'])
def disconnect():
    """
    Disconnect a social media connection
    Query params: provider (e.g., 'facebook')
    """
    provider = request.args.get('provider', 'facebook')

    try:
        conn = db.get_connection()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM social_connections WHERE provider = ?', (provider,))
        deleted = cursor.rowcount
        conn.commit()
        conn.close()

        if deleted > 0:
            return jsonify({
                'success': True,
                'message': f'{provider} connection removed'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'No connection found'
            })

    except Exception as e:
        print(f"❌ Error disconnecting: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/schedule-post', methods=['POST'])
def schedule_post():
    """
    Schedule a post for future publishing
    POST body: {
        "post_content": "content",
        "scheduled_time": "2025-11-24T10:00:00",
        "image_path": "/path/to/image.jpg",
        "video_path": "/path/to/video.mp4",
        "platform": "facebook"
    }
    """
    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    post_content = data.get('post_content')
    scheduled_time = data.get('scheduled_time')
    image_path = data.get('image_path')
    video_path = data.get('video_path')
    platform = data.get('platform', 'facebook')

    if not post_content or not scheduled_time:
        return jsonify({'error': 'post_content and scheduled_time are required'}), 400

    try:
        # Validate datetime format
        from datetime import datetime
        datetime.fromisoformat(scheduled_time)

        # Save to database
        post_id = db.schedule_post(
            post_content=post_content,
            scheduled_time=scheduled_time,
            image_path=image_path,
            video_path=video_path,
            platform=platform
        )

        return jsonify({
            'success': True,
            'message': 'Post scheduled successfully',
            'post_id': post_id,
            'scheduled_time': scheduled_time
        })

    except ValueError as e:
        return jsonify({'error': f'Invalid datetime format: {str(e)}'}), 400
    except Exception as e:
        print(f"❌ Error scheduling post: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/scheduled-posts', methods=['GET'])
def get_scheduled_posts():
    """Get all scheduled posts"""
    try:
        posts = db.get_all_scheduled_posts()
        return jsonify({
            'success': True,
            'posts': posts
        })
    except Exception as e:
        print(f"❌ Error getting scheduled posts: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/delete-scheduled-post/<int:post_id>', methods=['DELETE'])
def delete_scheduled_post(post_id):
    """Delete a scheduled post"""
    try:
        success = db.delete_scheduled_post(post_id)
        if success:
            return jsonify({
                'success': True,
                'message': f'Post {post_id} deleted successfully'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Failed to delete post'
            }), 500
    except Exception as e:
        print(f"❌ Error deleting scheduled post: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-from-url', methods=['POST'])
def generate_from_url():
    """
    Generate social media post from article URL
    POST body: {
        "url": "https://example.com/article",
        "language": "hu",  # hu, en, de, es, fr
        "style": "facebook"  # facebook, linkedin, instagram, twitter, tiktok, reels, shorts
    }
    """
    import requests
    from bs4 import BeautifulSoup

    data = request.get_json()

    if not data:
        return jsonify({'error': 'Request body is required'}), 400

    url = data.get('url')
    language = data.get('language', 'hu')
    style = data.get('style', 'facebook')

    if not url:
        return jsonify({'error': 'url is required'}), 400

    print(f"🔗 Generating post from URL: {url}")
    print(f"   Language: {language}, Style: {style}")

    try:
        # Fetch article content
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()

        # Parse HTML
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract title
        title = ''
        if soup.title:
            title = soup.title.string or ''
        og_title = soup.find('meta', property='og:title')
        if og_title and og_title.get('content'):
            title = og_title['content']

        # Extract description
        description = ''
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            description = meta_desc['content']
        og_desc = soup.find('meta', property='og:description')
        if og_desc and og_desc.get('content'):
            description = og_desc['content']

        # Extract article body text
        article_text = ''

        # Try to find article content
        article = soup.find('article')
        if article:
            paragraphs = article.find_all('p')
            article_text = ' '.join([p.get_text().strip() for p in paragraphs[:10]])
        else:
            # Fallback: get main content area
            main = soup.find('main') or soup.find('div', class_='content') or soup.find('div', class_='post-content')
            if main:
                paragraphs = main.find_all('p')
                article_text = ' '.join([p.get_text().strip() for p in paragraphs[:10]])
            else:
                # Last resort: get all paragraphs
                paragraphs = soup.find_all('p')
                article_text = ' '.join([p.get_text().strip() for p in paragraphs[:10]])

        # Truncate to reasonable length
        article_text = article_text[:2000] if article_text else description

        if not title and not article_text:
            return jsonify({'error': 'Could not extract content from URL'}), 400

        # Build prompt based on style
        language_names = {
            'hu': 'magyar',
            'en': 'English',
            'de': 'Deutsch',
            'es': 'Español',
            'fr': 'Français'
        }

        style_prompts = {
            'facebook': f"""Írj egy figyelemfelkeltő Facebook posztot erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A poszt legyen:
- Rövid (max 200 szó)
- Tartalmazzon 2-3 emojit
- Tartalmazzon 1-3 releváns hashtagot
- Tegyen fel egy gondolatébresztő kérdést a végén
- Legyen informatív de szórakoztató hangvételű""",

            'linkedin': f"""Írj egy professzionális LinkedIn posztot erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A poszt legyen:
- Szakmai hangvételű
- 150-250 szó
- Tartalmazzon kulcsfontosságú tanulságokat (bullet points)
- Végezzen gondolatébresztő kérdéssel vagy call-to-action-nel
- Releváns hashtagok a végén""",

            'instagram': f"""Írj egy Instagram caption-t erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A caption legyen:
- Rövid és ütős (max 150 szó)
- Tartalmazzon 3-5 emojit
- 5-10 releváns hashtag a végén
- Vizuális leírás javaslat a képhez""",

            'twitter': f"""Írj egy X/Twitter posztot erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A poszt legyen:
- Max 280 karakter
- Ütős és figyelemfelkeltő
- 1-2 hashtag
- Legyen rövid és tömör""",

            'tiktok': f"""Írj egy TikTok videó scriptet erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A script legyen:
- Hook az elején (első 3 másodperc)
- 30-60 másodperces videóhoz
- Dinamikus, gyors tempójú
- Közvetlen megszólítás (Te/Ti)
- CTA a végén (like, follow, comment)""",

            'reels': f"""Írj egy Facebook/Instagram Reels videó scriptet erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A script legyen:
- Erős hook az elején
- 15-30 másodperces videóhoz
- Vizuális tippekkel (mit mutassunk)
- Szórakoztató és informatív
- Emojikkal és dinamikával""",

            'shorts': f"""Írj egy YouTube Shorts videó scriptet erről a cikkről {language_names.get(language, 'magyar')} nyelven.
A script legyen:
- Figyelemfelkeltő nyitás
- Max 60 másodperces videóhoz
- Értékes tartalom gyorsan
- Subscribe CTA a végén
- Vizuális útmutatással"""
        }

        prompt = style_prompts.get(style, style_prompts['facebook'])

        full_prompt = f"""{prompt}

CIKK CÍME: {title}

CIKK TARTALMA:
{article_text}

---
Generáld le a posztot/scriptet a fenti utasítások alapján:"""

        print(f"   Generating with AI Provider: {AI_PROVIDER.upper()}")

        # Generate content using AI
        if AI_PROVIDER == 'google':
            content = google_ai_generator.generate_text(full_prompt)
        else:
            content = post_generator.generate_text(full_prompt)

        if not content:
            return jsonify({'error': 'AI generation failed'}), 500

        return jsonify({
            'success': True,
            'content': content,
            'title': title,
            'source_url': url,
            'style': style,
            'language': language
        })

    except requests.RequestException as e:
        print(f"❌ Error fetching URL: {e}")
        return jsonify({'error': f'Could not fetch URL: {str(e)}'}), 400
    except Exception as e:
        print(f"❌ Error generating from URL: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ============================================================================
# RAG STYLE LEARNING
# ============================================================================

@app.route('/api/style/upload', methods=['POST'])
def upload_style():
    """
    Upload text content to learn from as style reference.
    Accepts file upload or raw text.
    POST body: multipart/form-data with:
        - file: document file (.txt, .md, .docx, .pdf) OR
        - text: raw text content
        - source_name: identifier for this source (required)
        - style_name: style category (optional, default: "default")
    """
    from rag_store import get_rag_store

    source_name = request.form.get('source_name')
    style_name = request.form.get('style_name', 'default')

    if not source_name:
        return jsonify({'error': 'source_name is required'}), 400

    text_content = ""

    # Check for file upload
    if 'file' in request.files and request.files['file'].filename:
        file = request.files['file']
        filename = file.filename.lower()
        ext = filename.rsplit('.', 1)[-1] if '.' in filename else ''

        import tempfile
        import os as os_module

        temp_dir = tempfile.gettempdir()
        temp_path = os_module.path.join(temp_dir, f"style_upload_{file.filename}")
        file.save(temp_path)

        try:
            if ext == 'txt':
                with open(temp_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()

            elif ext == 'md':
                import markdown
                from bs4 import BeautifulSoup
                with open(temp_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                html = markdown.markdown(md_content)
                text_content = BeautifulSoup(html, 'html.parser').get_text()

            elif ext == 'docx':
                from docx import Document
                doc = Document(temp_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                text_content = '\n\n'.join(paragraphs)

            elif ext == 'pdf':
                import pdfplumber
                with pdfplumber.open(temp_path) as pdf:
                    pages_text = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            pages_text.append(text)
                    text_content = '\n\n'.join(pages_text)
            else:
                return jsonify({'error': f'Unsupported file type: {ext}'}), 400

            os_module.remove(temp_path)
        except Exception as e:
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500

    # Check for raw text
    elif 'text' in request.form:
        text_content = request.form.get('text', '')
    else:
        return jsonify({'error': 'Either file or text is required'}), 400

    if not text_content.strip():
        return jsonify({'error': 'No text content found'}), 400

    try:
        rag_store = get_rag_store()
        chunks_added = rag_store.add_style_sample(text_content, source_name, style_name)

        return jsonify({
            'success': True,
            'message': f'Style sample added successfully',
            'source_name': source_name,
            'style_name': style_name,
            'chunks_added': chunks_added,
            'text_length': len(text_content)
        })
    except Exception as e:
        print(f"❌ Error adding style sample: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/style/query', methods=['POST'])
def query_style():
    """
    Query the style store for similar text samples.
    POST body: {
        "query": "text to match",
        "n_results": 5,
        "source_filter": "optional source name",
        "style_filter": "optional style name"
    }
    """
    from rag_store import get_rag_store

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Request body required'}), 400

    query = data.get('query')
    if not query:
        return jsonify({'error': 'query is required'}), 400

    try:
        rag_store = get_rag_store()
        results = rag_store.query_style(
            query,
            n_results=data.get('n_results', 5),
            source_filter=data.get('source_filter'),
            style_filter=data.get('style_filter')
        )

        return jsonify({
            'success': True,
            'results': results
        })
    except Exception as e:
        print(f"❌ Error querying style: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/style/context', methods=['POST'])
def get_style_context():
    """
    Get style context for prompt augmentation.
    POST body: {
        "query": "topic or text to generate context for",
        "max_tokens": 1000
    }
    """
    from rag_store import get_rag_store

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Request body required'}), 400

    query = data.get('query')
    if not query:
        return jsonify({'error': 'query is required'}), 400

    try:
        rag_store = get_rag_store()
        context = rag_store.get_style_context(
            query,
            max_tokens=data.get('max_tokens', 1000)
        )

        return jsonify({
            'success': True,
            'context': context,
            'has_context': bool(context)
        })
    except Exception as e:
        print(f"❌ Error getting style context: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/style/sources', methods=['GET'])
def list_style_sources():
    """List all style sources in the store."""
    from rag_store import get_rag_store

    try:
        rag_store = get_rag_store()
        sources = rag_store.list_sources()

        return jsonify({
            'success': True,
            'sources': sources
        })
    except Exception as e:
        print(f"❌ Error listing sources: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/style/source/<source_name>', methods=['DELETE'])
def delete_style_source(source_name):
    """Delete all chunks from a specific source."""
    from rag_store import get_rag_store

    try:
        rag_store = get_rag_store()
        deleted_count = rag_store.delete_source(source_name)

        return jsonify({
            'success': True,
            'message': f'Deleted {deleted_count} chunks from {source_name}',
            'deleted_count': deleted_count
        })
    except Exception as e:
        print(f"❌ Error deleting source: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/style/stats', methods=['GET'])
def get_style_stats():
    """Get statistics about the style store."""
    from rag_store import get_rag_store

    try:
        rag_store = get_rag_store()
        stats = rag_store.get_stats()

        return jsonify({
            'success': True,
            'stats': stats
        })
    except Exception as e:
        print(f"❌ Error getting stats: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================================================
# DOCUMENT PARSING (Doksiból Poszt)
# ============================================================================

@app.route('/api/generate-from-doc', methods=['POST'])
def generate_from_doc():
    """
    Generate social media post from uploaded document
    Accepts: .docx, .pdf, .md, .txt files
    POST body: multipart/form-data with:
        - file: the document file
        - language: hu, en, de, es, fr (default: hu)
        - style: facebook, linkedin, instagram, twitter, tiktok, reels, shorts (default: facebook)
    """
    import tempfile
    import os as os_module

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    language = request.form.get('language', 'hu')
    style = request.form.get('style', 'facebook')

    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400

    # Get file extension
    filename = file.filename.lower()
    ext = filename.rsplit('.', 1)[-1] if '.' in filename else ''

    allowed_extensions = ['docx', 'pdf', 'md', 'txt']
    if ext not in allowed_extensions:
        return jsonify({'error': f'Unsupported file type. Allowed: {", ".join(allowed_extensions)}'}), 400

    print(f"📄 Processing document: {file.filename}")
    print(f"   Extension: {ext}, Language: {language}, Style: {style}")

    try:
        # Save file temporarily
        temp_dir = tempfile.gettempdir()
        temp_path = os_module.path.join(temp_dir, f"upload_{file.filename}")
        file.save(temp_path)

        # Extract text based on file type
        extracted_text = ""

        if ext == 'txt':
            with open(temp_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()

        elif ext == 'md':
            import markdown
            with open(temp_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            # Convert markdown to plain text (strip HTML)
            from bs4 import BeautifulSoup
            html = markdown.markdown(md_content)
            extracted_text = BeautifulSoup(html, 'html.parser').get_text()

        elif ext == 'docx':
            from docx import Document
            doc = Document(temp_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            extracted_text = '\n\n'.join(paragraphs)

        elif ext == 'pdf':
            import pdfplumber
            with pdfplumber.open(temp_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                extracted_text = '\n\n'.join(pages_text)

        # Clean up temp file
        os_module.remove(temp_path)

        if not extracted_text.strip():
            return jsonify({'error': 'Could not extract text from document'}), 400

        # Truncate if too long (max 10000 chars for AI)
        if len(extracted_text) > 10000:
            extracted_text = extracted_text[:10000] + "..."
            print(f"   ⚠️ Text truncated to 10000 characters")

        print(f"   ✅ Extracted {len(extracted_text)} characters")

        # Calculate SEO score using textstat
        seo_score = 0
        readability_info = {}
        try:
            import textstat
            # Set language for textstat
            if language == 'hu':
                textstat.set_lang('en')  # Hungarian not supported, use EN as fallback
            else:
                textstat.set_lang(language if language in ['en', 'de', 'es', 'fr'] else 'en')

            # Calculate various readability metrics
            flesch_score = textstat.flesch_reading_ease(extracted_text)
            grade_level = textstat.flesch_kincaid_grade(extracted_text)
            word_count = textstat.lexicon_count(extracted_text)
            sentence_count = textstat.sentence_count(extracted_text)

            # Calculate SEO score (0-100)
            # Higher flesch score = easier to read = better for social media
            seo_score = min(100, max(0, int(flesch_score)))

            readability_info = {
                'flesch_score': round(flesch_score, 1),
                'grade_level': round(grade_level, 1),
                'word_count': word_count,
                'sentence_count': sentence_count,
                'seo_score': seo_score
            }
            print(f"   📊 SEO Score: {seo_score}, Flesch: {flesch_score:.1f}")
        except Exception as e:
            print(f"   ⚠️ Could not calculate SEO score: {e}")
            seo_score = 50  # Default score

        # Build prompt based on style
        language_names = {
            'hu': 'magyar',
            'en': 'English',
            'de': 'Deutsch',
            'es': 'Español',
            'fr': 'Français'
        }

        style_prompts = {
            'facebook': f"""Írj egy figyelemfelkeltő Facebook posztot ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A poszt legyen:
- Rövid (max 200 szó)
- Tartalmazzon 2-3 emojit
- Tartalmazzon 1-3 releváns hashtagot
- Tegyen fel egy gondolatébresztő kérdést a végén
- Legyen informatív de szórakoztató hangvételű""",

            'linkedin': f"""Írj egy professzionális LinkedIn posztot ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A poszt legyen:
- Szakmai hangvételű
- 150-250 szó
- Tartalmazzon kulcsfontosságú tanulságokat (bullet points)
- Végezzen gondolatébresztő kérdéssel vagy call-to-action-nel
- Releváns hashtagok a végén""",

            'instagram': f"""Írj egy Instagram caption-t ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A caption legyen:
- Rövid és ütős (max 150 szó)
- Tartalmazzon 3-5 emojit
- 5-10 releváns hashtag a végén
- Vizuális leírás javaslat a képhez""",

            'twitter': f"""Írj egy X/Twitter posztot ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A poszt legyen:
- Max 280 karakter
- Ütős és figyelemfelkeltő
- 1-2 hashtag
- Legyen rövid és tömör""",

            'tiktok': f"""Írj egy TikTok videó scriptet ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A script legyen:
- Hook az elején (első 3 másodperc)
- 30-60 másodperces videóhoz
- Dinamikus, gyors tempójú
- Közvetlen megszólítás (Te/Ti)
- CTA a végén (like, follow, comment)""",

            'reels': f"""Írj egy Facebook/Instagram Reels videó scriptet ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A script legyen:
- Erős hook az elején
- 15-30 másodperces videóhoz
- Vizuális tippekkel (mit mutassunk)
- Szórakoztató és informatív
- Emojikkal és dinamikával""",

            'shorts': f"""Írj egy YouTube Shorts videó scriptet ebből a dokumentumból {language_names.get(language, 'magyar')} nyelven.
A script legyen:
- Figyelemfelkeltő nyitás
- Max 60 másodperces videóhoz
- Értékes tartalom gyorsan
- Subscribe CTA a végén
- Vizuális útmutatással"""
        }

        prompt = style_prompts.get(style, style_prompts['facebook'])

        full_prompt = f"""{prompt}

DOKUMENTUM TARTALMA:
{extracted_text}

---
Generáld le a posztot/scriptet a fenti utasítások alapján:"""

        print(f"   🤖 Generating with AI Provider: {AI_PROVIDER.upper()}")

        # Generate content using AI
        if AI_PROVIDER == 'google':
            content = google_ai_generator.generate_text(full_prompt)
        else:
            content = post_generator.generate_text(full_prompt)

        if not content:
            return jsonify({'error': 'AI generation failed'}), 500

        return jsonify({
            'success': True,
            'content': content,
            'filename': file.filename,
            'style': style,
            'language': language,
            'readability': readability_info,
            'text_length': len(extracted_text)
        })

    except Exception as e:
        print(f"❌ Error processing document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.errorhandler(404)
def not_found(e):
    """404 handler"""
    return jsonify({'error': 'Not found'}), 404


@app.errorhandler(500)
def server_error(e):
    """500 handler"""
    return jsonify({'error': 'Internal server error'}), 500


# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    port = int(os.getenv('PORT', 5000))
    debug = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'

    print(f"\n{'='*60}")
    print(f"🚀 TRENDMASTER STARTING")
    print(f"Port: {port}")
    print(f"Debug: {debug}")
    print(f"{'='*60}\n")

    app.run(host='0.0.0.0', port=port, debug=debug)
